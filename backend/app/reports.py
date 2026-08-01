"""Render a task's result (markdown-ish) into PDF and DOCX report bytes.

A lightweight markdown pass (headings, bullets, bold/code stripping) is enough
for task reports; no heavy HTML/CSS engine or system libraries required.
"""

import io
import re

from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from .models import Task


def _blocks(text: str):
    for raw in (text or "").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            yield ("blank", "")
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            yield (f"h{len(m.group(1))}", m.group(2))
            continue
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            yield ("bullet", re.sub(r"^\s*([-*]|\d+\.)\s+", "", line))
            continue
        yield ("para", line)


def _strip_md(s: str) -> str:
    s = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
    s = re.sub(r"\*(.*?)\*", r"\1", s)
    s = re.sub(r"`(.*?)`", r"\1", s)
    return s


def _meta_line(task: Task) -> str:
    when = task.completed_at or task.created_at
    model = task.model_used or task.model or "-"
    return (
        f"Status: {task.status}    Model: {model}    "
        f"Date: {when:%Y-%m-%d %H:%M UTC}" if when else f"Status: {task.status}"
    )


def build_docx(task: Task) -> bytes:
    doc = Document()
    doc.add_heading(task.title or "Task Report", level=0)
    doc.add_paragraph(_meta_line(task)).italic = True
    for typ, content in _blocks(task.result_text or "No report available for this task."):
        content = _strip_md(content)
        if typ.startswith("h"):
            doc.add_heading(content, level=min(int(typ[1]), 4))
        elif typ == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        elif typ == "para":
            doc.add_paragraph(content)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _latin1(s: str) -> str:
    # fpdf2 core fonts are latin-1; normalise common unicode then replace the rest.
    trans = {"—": "-", "–": "-", "•": "-", "“": '"', "”": '"', "‘": "'", "’": "'"}
    for k, v in trans.items():
        s = s.replace(k, v)
    return s.encode("latin-1", "replace").decode("latin-1")


def build_pdf(task: Task) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(True, margin=15)

    def line(text: str, h: float = 6) -> None:
        # new_x=LMARGIN resets the cursor to the left margin so the next
        # full-width cell has room (fpdf2 otherwise leaves x at the right edge).
        pdf.multi_cell(0, h, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "B", 16)
    line(_latin1(task.title or "Task Report"), 8)
    pdf.set_font("Helvetica", "I", 9)
    line(_latin1(_meta_line(task)), 5)
    pdf.ln(3)

    heading_size = {1: 15, 2: 13, 3: 12, 4: 11}
    for typ, content in _blocks(task.result_text or "No report available for this task."):
        content = _latin1(_strip_md(content))
        if typ == "blank":
            pdf.ln(2)
        elif typ.startswith("h"):
            pdf.set_font("Helvetica", "B", heading_size.get(int(typ[1]), 11))
            line(content, 7)
        elif typ == "bullet":
            pdf.set_font("Helvetica", "", 11)
            line(f"  -  {content}")
        else:
            pdf.set_font("Helvetica", "", 11)
            line(content)
    return bytes(pdf.output())
